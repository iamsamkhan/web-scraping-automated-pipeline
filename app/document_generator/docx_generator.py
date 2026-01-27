#
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List, Optional
import os
from pathlib import Path

class DocxGenerator:
    """Generate professional DOCX documents for academic papers"""
    
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
        
    def create_academic_paper(self, paper_data: Dict, output_path: str) -> str:
        """Create a DOCX academic paper"""
        
        # Create new document or use template
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
        else:
            doc = Document()
        
        # Apply styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading(paper_data['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author
        author = doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author.add_run(f"Author: {paper_data['author']}")
        author_run.bold = True
        author_run.font.size = Pt(12)
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {paper_data['date']}")
        
        doc.add_paragraph()  # Empty line
        
        # Add abstract section
        abstract_heading = doc.add_heading('Abstract', 1)
        abstract_para = doc.add_paragraph(paper_data['abstract'])
        abstract_para.style = 'BodyText'
        
        doc.add_page_break()
        
        # Add table of contents (simplified)
        toc_heading = doc.add_heading('Table of Contents', 1)
        
        sections = [
            '1. Introduction',
            '2. Literature Review',
            '3. Methodology',
            '4. Results',
            '5. Discussion',
            '6. Conclusion',
            'References',
            'Acknowledgments'
        ]
        
        for section in sections:
            doc.add_paragraph(section, style='TOC')
        
        doc.add_page_break()
        
        # Add main sections
        sections_data = {
            'Introduction': paper_data.get('introduction', ''),
            'Literature Review': paper_data.get('literature_review', ''),
            'Methodology': paper_data.get('methodology', ''),
            'Results': paper_data.get('results', ''),
            'Discussion': paper_data.get('discussion', ''),
            'Conclusion': paper_data.get('conclusion', '')
        }
        
        for section_title, section_content in sections_data.items():
            doc.add_heading(section_title, 1)
            doc.add_paragraph(section_content, style='BodyText')
            doc.add_paragraph()  # Empty line
        
        # Add references
        doc.add_heading('References', 1)
        for ref in paper_data.get('references', []):
            ref_para = doc.add_paragraph(ref, style='ListBullet')
        
        # Add acknowledgments
        doc.add_heading('Acknowledgments', 1)
        doc.add_paragraph(paper_data.get('acknowledgments', ''), style='BodyText')
        
        # Add contact info
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(paper_data.get('contact', ''))
        contact_run.italic = True
        
        # Add footer
        self._add_footer(doc, paper_data['author'])
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _setup_document_styles(self, doc: Document):
        """Setup document styles"""
        
        # Title style
        styles = doc.styles
        
        # Body text style
        if 'BodyText' not in styles:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Times New Roman'
            body_font.size = Pt(12)
        
        # TOC style
        if 'TOC' not in styles:
            toc_style = styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
            toc_font = toc_style.font
            toc_font.name = 'Arial'
            toc_font.size = Pt(11)
    
    def _add_footer(self, doc: Document, author: str):
        """Add footer to document"""
        
        section = doc.sections[0]
        footer = section.footer
        
        # Add page numbers
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add copyright
        copyright_para = footer.add_paragraph()
        copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_year = datetime.now().year
        copyright_run = copyright_para.add_run(f"© {current_year} {author}. Academic Paper. Confidential.")
        copyright_run.font.size = Pt(9)
        copyright_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def generate_personalized_paper(self, student_data: Dict, output_dir: str = "output/papers") -> Dict:
        """Generate personalized paper for a student"""
        
        # Create output directory if it doesn't exist
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Generate paper content
        from app.email_service.ai_paper_generator import AIPaperGenerator
        paper_gen = AIPaperGenerator(model_type="fallback")
        
        title = paper_gen.generate_paper_title(student_data.get('field', ''))
        abstract = paper_gen.generate_abstract(title, student_data['name'])
        
        paper_content = paper_gen.generate_paper_content(
            title=title,
            abstract=abstract,
            student_name=student_data['name']
        )
        
        # Create filename
        safe_name = "".join(c for c in student_data['name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Generate DOCX
        docx_path = self.create_academic_paper(paper_content, output_path)
        
        return {
            'student_name': student_data['name'],
            'student_email': student_data['email'],
            'paper_title': title,
            'abstract': abstract[:200] + "...",  # Preview
            'docx_path': docx_path,
            'file_size': os.path.getsize(docx_path) if os.path.exists(docx_path) else 0
        }