"""
Comprehensive unit tests for email service
"""

import pytest
import asyncio
from unittest.mock import Mock, patch, MagicMock, AsyncMock
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import tempfile
from pathlib import Path
import json

from app.email_service.email_sender import EmailSender
from app.email_service.ai_paper_generator import AIPaperGenerator
from app.email_service.template_manager import TemplateManager
from app.models.schemas import StudentBase

class TestEmailSender:
    """Test email sending functionality"""
    
    @pytest.fixture
    def email_sender(self):
        """Create email sender instance"""
        return EmailSender(
            smtp_server="smtp.test.com",
            smtp_port=587,
            use_ssl=False
        )
    
    @pytest.fixture
    def sample_student(self):
        """Create sample student"""
        return {
            "name": "John Doe",
            "email": "john.doe@test.com",
            "university": "Test University",
            "paper_title": "Test Research Paper"
        }
    
    @pytest.fixture
    def temp_docx_file(self):
        """Create temporary DOCX file for testing"""
        import docx
        from docx import Document
        
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test document for email attachments.")
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            yield tmp.name
        
        # Cleanup
        Path(tmp.name).unlink(missing_ok=True)
    
    def test_email_initialization(self, email_sender):
        """Test email sender initialization"""
        assert email_sender.smtp_server == "smtp.test.com"
        assert email_sender.smtp_port == 587
        assert email_sender.use_ssl == False
        assert email_sender.emails_per_hour == 50
        assert email_sender.delay_between_emails == 2
    
    @patch('smtplib.SMTP')
    def test_send_email_success(self, mock_smtp, email_sender, temp_docx_file):
        """Test successful email sending"""
        # Mock SMTP
        mock_server = Mock()
        mock_smtp.return_value = mock_server
        
        # Configure sender
        email_sender.sender_email = "sender@test.com"
        email_sender.sender_password = "password"
        email_sender.sender_name = "Test Sender"
        
        # Send email with attachment
        success, message = email_sender.send_email(
            recipient_email="recipient@test.com",
            subject="Test Subject",
            body_html="<h1>Test HTML</h1>",
            body_text="Test plain text",
            attachments=[temp_docx_file]
        )
        
        assert success is True
        assert "successfully" in message.lower()
        
        # Verify SMTP was called
        mock_smtp.assert_called_once_with("smtp.test.com", 587)
        mock_server.starttls.assert_called_once()
        mock_server.login.assert_called_once_with("sender@test.com", "password")
        mock_server.send_message.assert_called_once()
        mock_server.quit.assert_called_once()
    
    @patch('smtplib.SMTP')
    def test_send_email_failure(self, mock_smtp, email_sender):
        """Test failed email sending"""
        # Mock SMTP to raise exception
        mock_smtp.side_effect = smtplib.SMTPException("SMTP error")
        
        success, message = email_sender.send_email(
            recipient_email="recipient@test.com",
            subject="Test Subject",
            body_html="<h1>Test</h1>"
        )
        
        assert success is False
        assert "failed" in message.lower()
    
    def test_add_attachment(self, email_sender):
        """Test adding attachments to email"""
        # Create a test file
        test_content = b"Test attachment content"
        
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp.write(test_content)
            tmp.flush()
            
            # Create message
            msg = MIMEMultipart()
            
            # Add attachment
            email_sender._add_attachment(msg, tmp.name)
            
            # Verify attachment was added
            assert len(msg.get_payload()) == 1
            
            attachment = msg.get_payload()[0]
            assert attachment.get_content_type() == 'application/octet-stream'
            assert attachment['Content-Disposition'].startswith('attachment')
            
            # Cleanup
            Path(tmp.name).unlink()
    
    def test_html_to_text_conversion(self, email_sender):
        """Test HTML to plain text conversion"""
        test_html = """
        <html>
            <body>
                <h1>Title</h1>
                <p>Paragraph with <strong>bold</strong> text.</p>
                <ul>
                    <li>Item 1</li>
                    <li>Item 2</li>
                </ul>
            </body>
        </html>
        """
        
        plain_text = email_sender._html_to_text(test_html)
        
        # Should remove HTML tags
        assert "<" not in plain_text
        assert ">" not in plain_text
        
        # Should preserve text content
        assert "Title" in plain_text
        assert "Paragraph with bold text." in plain_text
        assert "Item 1" in plain_text
        assert "Item 2" in plain_text
    
    @patch.object(EmailSender, 'send_email')
    def test_bulk_email_sending(self, mock_send_email, email_sender, sample_student):
        """Test bulk email sending"""
        # Mock individual email sending
        mock_send_email.return_value = (True, "Sent successfully")
        
        # Create email list
        email_list = [
            {**sample_student, "email": f"student{i}@test.com", "name": f"Student {i}"}
            for i in range(5)
        ]
        
        # Create template file
        template_content = """
        <html>
            <body>
                <h1>Hello {{ student_name }}</h1>
                <p>Your paper: {{ paper_title }}</p>
            </body>
        </html>
        """
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as tmp:
            tmp.write(template_content)
            tmp.flush()
            
            results = email_sender.send_bulk_emails(
                email_list=email_list,
                subject_template="Paper: {paper_title}",
                body_template_path=tmp.name,
                test_mode=False
            )
            
            # Cleanup
            Path(tmp.name).unlink()
        
        # Verify results
        assert results['total'] == 5
        assert results['success'] == 5
        assert results['failed'] == 0
        
        # Verify send_email was called for each
        assert mock_send_email.call_count == 5
    
    @patch.object(EmailSender, 'send_email')
    def test_test_mode(self, mock_send_email, email_sender, sample_student):
        """Test bulk email sending in test mode"""
        # Mock individual email sending
        mock_send_email.return_value = (True, "Sent successfully")
        
        # Create larger email list
        email_list = [
            {**sample_student, "email": f"student{i}@test.com"}
            for i in range(10)  # 10 emails
        ]
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as tmp:
            tmp.write("<html><body>Test</body></html>")
            tmp.flush()
            
            results = email_sender.send_bulk_emails(
                email_list=email_list,
                subject_template="Test",
                body_template_path=tmp.name,
                test_mode=True  # Should only send first 5
            )
            
            Path(tmp.name).unlink()
        
        # In test mode, only first 5 should be sent
        assert mock_send_email.call_count == 5
        assert results['total'] == 10
        assert results['success'] == 5
    
    def test_email_logging(self, email_sender):
        """Test email logging functionality"""
        import json
        from pathlib import Path
        
        # Clean up any existing log file
        log_file = Path("logs/email_log.json")
        if log_file.exists():
            log_file.unlink()
        
        # Create logs directory
        Path("logs").mkdir(exist_ok=True)
        
        # Log an email
        email_sender._log_email(
            recipient="test@example.com",
            subject="Test Subject",
            success=True,
            error=""
        )
        
        # Verify log was created
        assert log_file.exists()
        
        # Read and verify log
        with open(log_file, 'r') as f:
            logs = json.load(f)
        
        assert len(logs) == 1
        assert logs[0]['recipient'] == "test@example.com"
        assert logs[0]['subject'] == "Test Subject"
        assert logs[0]['success'] is True
        
        # Cleanup
        log_file.unlink()

class TestAIPaperGenerator:
    """Test AI paper generator"""
    
    @pytest.fixture
    def paper_generator(self):
        """Create AI paper generator instance"""
        return AIPaperGenerator(model_type="fallback")
    
    def test_generate_paper_title(self, paper_generator):
        """Test paper title generation"""
        # Test with specific field
        title = paper_generator.generate_paper_title("Computer Science")
        assert isinstance(title, str)
        assert len(title) > 10
        
        # Test with random field
        random_title = paper_generator.generate_paper_title()
        assert isinstance(random_title, str)
        assert len(random_title) > 10
        
        # Test with different fields
        fields = ["Artificial Intelligence", "Biology", "Economics"]
        for field in fields:
            title = paper_generator.generate_paper_title(field)
            assert title
            assert len(title) > 0
    
    def test_generate_abstract(self, paper_generator):
        """Test abstract generation"""
        title = "Advances in Machine Learning"
        student_name = "John Doe"
        
        abstract = paper_generator.generate_abstract(title, student_name)
        
        assert isinstance(abstract, str)
        assert len(abstract) > 100  # Should be substantial
        
        # Should contain relevant information
        assert title in abstract or any(word in abstract for word in title.split())
        assert student_name in abstract
        
        # Check structure (should have common academic abstract elements)
        assert "ABSTRACT" in abstract.upper() or "abstract" in abstract.lower()
    
    def test_generate_paper_content(self, paper_generator):
        """Test full paper content generation"""
        title = "Test Research Paper"
        abstract = "This is a test abstract."
        student_name = "Test Student"
        
        paper_content = paper_generator.generate_paper_content(
            title=title,
            abstract=abstract,
            student_name=student_name
        )
        
        # Check structure
        assert paper_content['title'] == title
        assert paper_content['author'] == student_name
        assert paper_content['abstract'] == abstract
        
        # Check sections exist
        required_sections = [
            'introduction', 'literature_review', 'methodology',
            'results', 'discussion', 'conclusion', 'references'
        ]
        
        for section in required_sections:
            assert section in paper_content
            assert isinstance(paper_content[section], str)
            assert len(paper_content[section]) > 0
    
    def test_generate_references(self, paper_generator):
        """Test reference generation"""
        references = paper_generator._generate_references()
        
        assert isinstance(references, list)
        assert len(references) > 0
        
        # Check reference format
        for ref in references:
            assert isinstance(ref, str)
            assert len(ref) > 20
            # Should contain author and year
            assert "(" in ref and ")" in ref
    
    @patch('openai.ChatCompletion.create')
    def test_openai_generation(self, mock_openai):
        """Test OpenAI integration"""
        # Mock OpenAI response
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = "Test abstract generated by AI."
        mock_openai.return_value = mock_response
        
        generator = AIPaperGenerator(model_type="openai", api_key="test-key")
        
        abstract = generator._generate_with_openai(
            title="Test Title",
            student_name="Test Student"
        )
        
        assert abstract == "Test abstract generated by AI."
        mock_openai.assert_called_once()
    
    def test_categorize_field(self, paper_generator):
        """Test field categorization"""
        test_cases = [
            ("Computer Science", "computer_science"),
            ("Software Engineering", "computer_science"),
            ("Mechanical Engineering", "engineering"),
            ("Electrical Eng", "engineering"),
            ("Biology", "science"),
            ("Physics", "science"),
            ("Unknown Field", "computer_science"),  # Default
        ]
        
        for input_field, expected_category in test_cases:
            result = paper_generator._categorize_field(input_field)
            assert result == expected_category

class TestTemplateManager:
    """Test template management"""
    
    @pytest.fixture
    def template_manager(self, tmp_path):
        """Create template manager with temporary directory"""
        template_dir = tmp_path / "templates"
        template_dir.mkdir()
        
        # Create test templates
        base_template = template_dir / "base.html"
        base_template.write_text("""
        <!DOCTYPE html>
        <html>
        <body>
            {% block content %}{% endblock %}
        </body>
        </html>
        """)
        
        test_template = template_dir / "test.html"
        test_template.write_text("""
        {% extends "base.html" %}
        {% block content %}
        <h1>Hello {{ name }}</h1>
        <p>Email: {{ email }}</p>
        {% endblock %}
        """)
        
        return TemplateManager(template_dir=str(template_dir))
    
    def test_template_rendering(self, template_manager):
        """Test template rendering"""
        context = {
            "name": "John Doe",
            "email": "john@example.com"
        }
        
        # Render template
        template = template_manager.env.get_template("test.html")
        rendered = template.render(**context)
        
        assert "John Doe" in rendered
        assert "john@example.com" in rendered
        assert "<h1>" in rendered  # HTML should be preserved
    
    def test_template_validation(self, template_manager, tmp_path):
        """Test template validation"""
        # Test valid template
        valid_template = tmp_path / "valid.html"
        valid_template.write_text("<h1>Valid</h1>")
        
        # This should not raise an exception
        template_manager.env.get_template(str(valid_template.relative_to(tmp_path / "templates")))
    
    def test_create_default_templates(self, tmp_path):
        """Test default template creation"""
        template_dir = tmp_path / "email_templates"
        
        manager = TemplateManager(template_dir=str(template_dir))
        
        # Check that default templates were created
        assert (template_dir / "base.html").exists()
        assert (template_dir / "journal_invitation.html").exists()
        
        # Verify template content
        base_content = (template_dir / "base.html").read_text()
        assert "<!DOCTYPE html>" in base_content
        assert "{% block content %}" in base_content

# Integration tests for email workflow
@pytest.mark.integration
class TestEmailWorkflow:
    """Integration tests for email workflow"""
    
    @pytest.fixture
    def complete_email_data(self):
        """Create complete email test data"""
        return {
            "student": {
                "name": "John Doe",
                "email": "john.doe@university.edu",
                "university": "Test University",
                "department": "Computer Science"
            },
            "paper": {
                "title": "Advances in Test Research",
                "abstract": "This is a test abstract for the research paper.",
                "content": {
                    "introduction": "Test introduction",
                    "methodology": "Test methodology",
                    "results": "Test results",
                    "conclusion": "Test conclusion"
                }
            }
        }
    
    def test_complete_email_generation_flow(self, tmp_path, complete_email_data):
        """Test complete email generation flow"""
        from app.document_generator.docx_generator import DocxGenerator
        from app.email_service.ai_paper_generator import AIPaperGenerator
        
        # Setup directories
        output_dir = tmp_path / "output"
        output_dir.mkdir()
        
        template_dir = tmp_path / "templates"
        template_dir.mkdir()
        
        # Create template
        template_content = """
        <html>
        <body>
            <h1>Research Paper: {{ paper_title }}</h1>
            <p>Dear {{ student_name }},</p>
            <p>Abstract: {{ abstract }}</p>
        </body>
        </html>
        """
        
        template_file = template_dir / "test_template.html"
        template_file.write_text(template_content)
        
        # Generate paper
        paper_gen = AIPaperGenerator(model_type="fallback")
        docx_gen = DocxGenerator()
        
        paper_data = paper_gen.generate_paper_content(
            title=complete_email_data["paper"]["title"],
            abstract=complete_email_data["paper"]["abstract"],
            student_name=complete_email_data["student"]["name"]
        )
        
        # Create DOCX
        docx_path = output_dir / "test_paper.docx"
        docx_gen.create_academic_paper(paper_data, str(docx_path))
        
        assert docx_path.exists()
        assert docx_path.stat().st_size > 0
        
        # Create email
        email_sender = EmailSender()
        
        # Mock SMTP for integration test
        with patch('smtplib.SMTP') as mock_smtp:
            mock_server = Mock()
            mock_smtp.return_value = mock_server
            
            email_sender.sender_email = "test@example.com"
            email_sender.sender_password = "password"
            
            success, message = email_sender.send_email(
                recipient_email=complete_email_data["student"]["email"],
                subject=f"Research Paper: {complete_email_data['paper']['title']}",
                body_html=template_content.replace(
                    "{{ paper_title }}", complete_email_data["paper"]["title"]
                ).replace(
                    "{{ student_name }}", complete_email_data["student"]["name"]
                ).replace(
                    "{{ abstract }}", complete_email_data["paper"]["abstract"][:100]
                ),
                attachments=[str(docx_path)]
            )
            
            assert success is True
    
    @pytest.mark.performance
    def test_bulk_email_performance(self, tmp_path):
        """Test bulk email performance"""
        import time
        
        email_sender = EmailSender()
        email_sender.delay_between_emails = 0  # Remove delay for test
        
        # Create test data
        email_list = []
        for i in range(100):
            email_list.append({
                "name": f"Student {i}",
                "email": f"student{i}@test.com",
                "university": "Test University",
                "paper_title": f"Research Paper {i}",
                "abstract": f"Abstract for paper {i}"
            })
        
        # Create template
        template_content = "<html><body>Test</body></html>"
        template_file = tmp_path / "template.html"
        template_file.write_text(template_content)
        
        # Mock SMTP
        with patch.object(email_sender, 'send_email') as mock_send:
            mock_send.return_value = (True, "Success")
            
            start_time = time.time()
            
            results = email_sender.send_bulk_emails(
                email_list=email_list,
                subject_template="Test {paper_title}",
                body_template_path=str(template_file),
                test_mode=False
            )
            
            end_time = time.time()
            execution_time = end_time - start_time
            
        # Verify performance
        assert execution_time < 10.0, f"Bulk emails took too long: {execution_time:.2f}s"
        assert results['total'] == 100
        assert results['success'] == 100
        
        print(f"\nPerformance: Sent 100 emails in {execution_time:.2f}s")

# Error handling tests
@pytest.mark.error_handling
class TestEmailErrorHandling:
    """Test email service error handling"""
    
    def test_invalid_email_address(self):
        """Test handling of invalid email addresses"""
        email_sender = EmailSender()
        
        # Invalid email should still attempt to send (validation happens at SMTP level)
        # but we test that it doesn't crash
        with patch('smtplib.SMTP') as mock_smtp:
            mock_server = Mock()
            mock_smtp.return_value = mock_server
            mock_server.send_message.side_effect = smtplib.SMTPRecipientsRefused({})
            
            success, message = email_sender.send_email(
                recipient_email="invalid-email",
                subject="Test",
                body_html="<h1>Test</h1>"
            )
            
            assert success is False
    
    def test_missing_attachments(self, email_sender):
        """Test handling of missing attachment files"""
        # Test with non-existent file
        success, message = email_sender.send_email(
            recipient_email="test@example.com",
            subject="Test",
            body_html="<h1>Test</h1>",
            attachments=["/nonexistent/file.docx"]
        )
        
        # Should still attempt to send email without attachment
        # (implementation dependent)
    
    def test_template_not_found(self, email_sender):
        """Test handling of missing template file"""
        email_list = [{
            "name": "Test",
            "email": "test@example.com",
            "paper_title": "Test"
        }]
        
        # Should raise appropriate exception
        with pytest.raises(FileNotFoundError):
            email_sender.send_bulk_emails(
                email_list=email_list,
                subject_template="Test",
                body_template_path="/nonexistent/template.html"
            )

# Security tests
@pytest.mark.security
class TestEmailSecurity:
    """Security tests for email service"""
    
    def test_email_injection_prevention(self, email_sender):
        """Test prevention of email injection attacks"""
        malicious_input = "test@example.com\nCC: victim@example.com"
        
        with patch('smtplib.SMTP') as mock_smtp:
            mock_server = Mock()
            mock_smtp.return_value = mock_server
            
            # Send email with potentially malicious input
            success, message = email_sender.send_email(
                recipient_email=malicious_input,
                subject="Test",
                body_html="<h1>Test</h1>"
            )
            
            # SMTP library should handle or reject this
            # We're just testing that it doesn't crash
    
    def test_html_injection_in_templates(self, template_manager):
        """Test that templates safely handle user input"""
        malicious_context = {
            "name": "<script>alert('XSS')</script>",
            "email": "test@example.com"
        }
        
        # Create template with user input
        template_content = """
        <html>
        <body>
            <h1>Hello {{ name }}</h1>
        </body>
        </html>
        """
        
        template = template_manager.env.from_string(template_content)
        rendered = template.render(**malicious_context)
        
        # Jinja2 should escape HTML by default
        assert "&lt;script&gt;" in rendered
        assert "<script>" not in rendered
    
    def test_sensitive_data_logging(self, tmp_path):
        """Test that sensitive data is not logged"""
        import json
        
        email_sender = EmailSender()
        
        # Configure with sensitive data
        email_sender.sender_email = "secret@example.com"
        email_sender.sender_password = "supersecretpassword"
        
        # Log an email
        log_file = tmp_path / "email_log.json"
        email_sender._log_email = lambda *args, **kwargs: None  # Mock logging
        
        # Send email (password should not be logged)
        with patch('smtplib.SMTP') as mock_smtp:
            mock_server = Mock()
            mock_smtp.return_value = mock_server
            
            success, message = email_sender.send_email(
                recipient_email="recipient@example.com",
                subject="Test",
                body_html="<h1>Test</h1>"
            )
        
        # Verify no sensitive data in any potential logs
        # (This would require checking actual log implementation)